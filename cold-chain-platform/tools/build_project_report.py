from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor


OUT = r"D:\zuoyiqing\cold-chain-platform\docs\生鲜物流协同优化系统项目展示材料.docx"


def set_run(run, size=10.5, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def paragraph(doc, text="", style=None, size=10.5, bold=False):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        set_run(run, size=16 if level == 1 else 13, bold=True, color=(22, 78, 99))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run(r, 9.5, True, (255, 255, 255))
        cell._tc.get_or_add_tcPr().append(parse_shd("164e63"))
        if widths:
            cell.width = widths[i]
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cells[i].width = widths[i]
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run(r, 9)
    doc.add_paragraph()
    return t


def parse_shd(fill):
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    return parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), fill))


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run(r, 10)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("生鲜物流协同优化系统\n项目展示与论文支撑材料")
set_run(r, 22, True, (15, 76, 92))
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("需求分析 · 系统架构 · 数据库设计 · 算法流程 · 测试与创新点")
set_run(r, 11, False, (80, 80, 80))
paragraph(doc, "适用场景：课程设计、毕业设计、项目答辩、软件验收说明。", size=10)
doc.add_page_break()

heading(doc, "1. 需求分析")
paragraph(doc, "系统面向生鲜冷链配送场景，核心目标是在客户时间窗、车辆容量、周转箱库存和装车顺序等约束下，生成成本可控、准时率高、可执行的配送方案。")
table(
    doc,
    ["角色", "主要诉求", "系统能力"],
    [
        ["调度员", "快速生成可执行路线，识别迟到风险", "GA-2-opt 求解、候选方案对比、异常重优化"],
        ["仓库员", "按卸货顺序完成装车，避免翻箱", "三维装箱、容量校验、周转箱库存管理"],
        ["司机", "清晰查看任务、确认到达、上报异常", "司机任务、派单状态流、送达回写"],
        ["管理员", "保障数据、权限和系统运行", "鉴权、角色权限、健康诊断、历史归档"],
    ],
    [Cm(2.4), Cm(5.2), Cm(8.0)],
)

heading(doc, "2. 系统架构")
bullets(doc, [
    "前端：Vue 3、TypeScript、Element Plus、Three.js，采用路由懒加载和模块化 API 封装。",
    "后端：FastAPI、SQLAlchemy、Alembic、PostgreSQL，提供订单、车辆、箱型、优化、派单、异常和诊断接口。",
    "算法层：GA-2-opt 负责路径与时间窗求解，规则装箱器负责三维装车坐标与容量可行性校验。",
    "桌面端：Electron 负责一键启动 Docker Compose 服务并打开本地系统界面。",
])
table(
    doc,
    ["层次", "组件", "职责"],
    [
        ["表现层", "Dashboard / Orders / Optimization / Packing / Driver", "运营总览、订单复核、地图路线、装箱和司机执行"],
        ["服务层", "FastAPI Routers", "鉴权、业务接口、异常处理、历史方案查询"],
        ["数据层", "PostgreSQL + Alembic", "订单、车辆、箱型、优化任务、派单、送达和异常记录"],
        ["算法层", "GA-2-opt + Packing Solver", "时间窗路径规划、容量约束、装车顺序和风险指标"],
    ],
    [Cm(2.5), Cm(5), Cm(8.2)],
)

heading(doc, "3. 数据库 ER 设计")
table(
    doc,
    ["实体", "关键字段", "关系说明"],
    [
        ["users", "username, role, status", "控制登录身份和角色权限"],
        ["orders", "order_no, destination, time window, geocode_status", "关联箱型、配送记录、派单和异常"],
        ["vehicles", "plate_no, capacity, status", "参与优化与派单执行"],
        ["box_types", "code, dimensions, stock_quantity", "支持逐箱型库存校验"],
        ["optimization_tasks / solutions", "objective, metrics, payload", "保存历史方案和对比指标"],
        ["dispatch_assignments", "order_id, vehicle_id, sequence, status", "承接派单状态流"],
        ["delivery_records", "actual_arrival, status", "司机送达回写"],
        ["operation_exceptions", "exception_type, status, resolution", "支持异常登记和重优化"],
    ],
    [Cm(4), Cm(5.5), Cm(6.3)],
)

heading(doc, "4. 算法流程")
bullets(doc, [
    "数据预处理：检查订单时间窗、经纬度复核状态、箱型和周转箱库存。",
    "距离与时间矩阵：基于配送中心和订单坐标计算行驶距离、行驶时间和服务时间。",
    "约束建模：建立时间维度、重量维度、体积维度、车辆固定成本和迟到软惩罚。",
    "求解与评估：使用 GA-2-opt 生成 ???? 方案，同时保留启发式候选作为方案对照。",
    "装箱校验：根据车辆尺寸、箱型尺寸和卸货顺序生成三维坐标，标记超容和装载风险。",
])

heading(doc, "5. 测试用例")
table(
    doc,
    ["编号", "测试点", "预期结果"],
    [
        ["T01", "周转箱逐箱型库存不足", "接口返回 400，并指出具体箱型缺口"],
        ["T02", "兜底定位订单参与优化", "系统阻止优化并提示人工复核坐标"],
        ["T03", "车辆超重或超容", "方案标记为不可行，不进入可执行派单"],
        ["T04", "装箱数量超过车辆格位", "返回 overflow_count，不重叠显示箱子"],
        ["T05", "司机确认送达", "写入 delivery_records，并同步订单/派单状态"],
        ["T06", "异常事件触发重优化", "异常落库，调度端可重新生成方案"],
    ],
    [Cm(2), Cm(6), Cm(7.8)],
)

heading(doc, "6. 性能与构建结果")
bullets(doc, [
    "前端入口包经过路由懒加载和 manualChunks 拆分，入口 JS 约 49KB，Three.js、xlsx、Element Plus 独立加载。",
    "后端核心测试覆盖库存、装箱和容量可行性，pytest 已通过。",
    "桌面端使用 Electron Builder 重新打包，快捷方式指向的新构建可加载更新后的代码。",
])

heading(doc, "7. 界面截图建议")
table(
    doc,
    ["截图位置", "展示重点"],
    [
        ["运营总览", "订单、车辆、成本、准时率、历史方案和异常事件"],
        ["订单管理", "地址复核、箱型、时间窗和订单状态"],
        ["路径优化", "路线地图、迟到点、候选方案指标"],
        ["三维装箱", "箱体颜色、装载顺序、容量和超容提示"],
        ["司机任务", "车辆任务、预计到达、确认送达"],
    ],
    [Cm(4), Cm(11.8)],
)

heading(doc, "8. 创新点")
bullets(doc, [
    "将 ???? 路径优化、三维装箱和司机执行闭环放在同一个系统中，不只停留在单一算法演示。",
    "把地址复核、周转箱库存和装载容量作为优化前置约束，提高方案可信度。",
    "保留历史方案对比，使成本、准时率、碳排和车辆数可以用于复盘和答辩说明。",
    "通过桌面端封装 Docker 启动流程，降低非技术用户运行系统的门槛。",
])

heading(doc, "9. 后续展望")
bullets(doc, [
    "接入真实地图路网矩阵和实时交通，替换经纬度估算距离。",
    "增加司机定位、电子签收、温湿度 IoT 数据和客户通知。",
    "扩展多仓、多温区、多批次波次配送和动态加单。",
])

doc.save(OUT)
print(OUT)
